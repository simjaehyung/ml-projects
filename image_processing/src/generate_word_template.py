import sys
import subprocess

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not found. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

# Main Title
head = document.add_heading('팀 활동 보고서', level=0)
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_paragraph('\n')

# Section 1: Basic Info
p = document.add_paragraph()
run = p.add_run('■ 기본 정보')
run.bold = True
run.font.size = Pt(12)

table1 = document.add_table(rows=4, cols=2)
table1.style = 'Table Grid'

data1 = [
    ('팀 명', ''),
    ('작성자', ''),
    ('제출일자', ''),
    ('참여자', '')
]

for i, (hdr, val) in enumerate(data1):
    table1.rows[i].cells[0].text = hdr
    table1.rows[i].cells[1].text = val

document.add_paragraph('\n')

# Section 2: Activity Details
p2 = document.add_paragraph()
run2 = p2.add_run('■ 활동 개요 및 내용')
run2.bold = True
run2.font.size = Pt(12)

table2 = document.add_table(rows=5, cols=2)
table2.style = 'Table Grid'

data2 = [
    ('활동 주제', ''),
    ('활동 기간', ''),
    ('활동 목적', '\n\n'),
    ('주요 활동 내용', '\n\n\n\n\n\n'),
    ('역할 분담', '\n\n\n')
]

for i, (hdr, val) in enumerate(data2):
    table2.rows[i].cells[0].text = hdr
    table2.rows[i].cells[1].text = val

document.add_paragraph('\n')

# Section 3: Results & Future Plans
p3 = document.add_paragraph()
run3 = p3.add_run('■ 결과 및 향후 계획')
run3.bold = True
run3.font.size = Pt(12)

table3 = document.add_table(rows=3, cols=2)
table3.style = 'Table Grid'

data3 = [
    ('결과 및 성과 (배운 점)', '\n\n\n\n'),
    ('문제점 및 해결 방안', '\n\n\n'),
    ('향후 계획', '\n\n')
]

for i, (hdr, val) in enumerate(data3):
    table3.rows[i].cells[0].text = hdr
    table3.rows[i].cells[1].text = val

document.save('팀활동보고서_양식.docx')
print("success")
